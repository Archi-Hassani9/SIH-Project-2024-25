import streamlit as st
import pandas as pd
from scholarly import scholarly, MaxTriesExceededException
from docx import Document
import bibtexparser
import io
import time

class FileHandler:
    def __init__(self, file):
        self.file = file
        self.fn = file.name

    def read_file(self):
        if self.fn.endswith('.bib'):
            content = self.file.read().decode('utf-8')
            bib_db = bibtexparser.load(io.StringIO(content))
            return pd.DataFrame(bib_db.entries)
        elif self.fn.endswith('.xlsx'):
            return pd.read_excel(self.file)
        else:
            st.error("Unsupported file format. Please upload a .bib or .xlsx file.")
            return None

class PublicationFetcher:
    def __init__(self, author_name):
        self.author_name = author_name

    def get_publications(self):
        query = scholarly.search_author(self.author_name)
        for _ in range(3):
            try:
                author = scholarly.fill(next(query))
                break
            except (StopIteration, MaxTriesExceededException):
                time.sleep(5)
        else:
            st.error("Failed to fetch data from Google Scholar. Try again later.")
            return pd.DataFrame()

        pubs = [
            {
                'title': pub.get('bib', {}).get('title', ''),
                'year': pub.get('bib', {}).get('pub_year', ''),
                'journal': pub.get('bib', {}).get('venue', ''),
                'authors': pub.get('bib', {}).get('author', '')
            }
            for pub in author['publications']
        ]
        return pd.DataFrame(pubs)

class DataProcessor:
    @staticmethod
    def filter_by_year(data, start, end):
        data['year'] = pd.to_numeric(data['year'], errors='coerce')
        return data[(data['year'] >= start) & (data['year'] <= end)]

class DataSaver:
    @staticmethod
    def save_to_excel(data, filename):
        data.to_excel(filename, index=False)
        st.success(f"Data saved to {filename}")

    @staticmethod
    def save_to_word(data, filename):
        doc = Document()
        doc.add_heading('Publication Summary', 0)
        for _, row in data.iterrows():
            doc.add_paragraph(f"{row['year']}: {row['title']} ({row['journal']})")
        doc.save(filename)
        st.success(f"Data saved to {filename}")

def main():
    st.title("Faculty Publication Summary Tool")
    df = st.sidebar.file_uploader("Upload a BibTeX or Excel file", type=["bib", "xlsx"])

    if df:
        file_handler = FileHandler(df)
        data = file_handler.read_file()
        if data is not None:
            st.write("### Uploaded Data", data.head())
            if 'authorName' not in st.session_state:
                st.session_state.authorName = "Jane Doe"
            author_name = st.text_input("Enter Author Name", st.session_state.authorName)

            search_opt = st.radio(
                "Choose search option:",
                ('Search within the uploaded dataset', 'Search universally using Google Scholar')
            )

            if st.button("Get Publications"):
                st.session_state.authorName = author_name
                if search_opt == 'Search within the uploaded dataset':
                    publications = data[data['author'].str.contains(author_name, case=False, na=False)]
                    if not publications.empty:
                        st.session_state.publications = publications
                        st.write(f"### Publications for {author_name} (from uploaded dataset)", publications)
                    else:
                        st.warning(f"No publications found for {author_name} in the uploaded dataset")
                else:
                    fetcher = PublicationFetcher(author_name)
                    publications = fetcher.get_publications()
                    if not publications.empty:
                        st.session_state.publications = publications
                        st.write(f"### Publications for {author_name} (from Google Scholar)", publications)
                    else:
                        st.warning(f"No publications found for {author_name}")

            if 'publications' in st.session_state:
                year1 = st.slider("Start Year", 1900, 2024, 2015)
                year2 = st.slider("End Year", 1900, 2024, 2020)
                filtered_data = DataProcessor.filter_by_year(st.session_state.publications, year1, year2)
                st.write(f"### Filtered Publications ({year1}-{year2})", filtered_data)

                if st.button("Save to Excel"):
                    DataSaver.save_to_excel(filtered_data, 'publication_summary.xlsx')
                if st.button("Save to Word"):
                    DataSaver.save_to_word(filtered_data, 'publication_summary.docx')

if __name__ == "__main__":
    main()
