import streamlit as st
import pandas as pd
from scholarly import scholarly, MaxTriesExceededException
from docx import Document
import bibtexparser
from io import BytesIO
import time

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# Defining a function called readdf that takes the user uploaded file as an argument and returns a readable file (a dataframe) 

# The st.cache_data decorator puts the user uploaded file into a kind of temperory memory so that everytime the user wishes to interact with it, the page does not reload
@st.cache_data
def readdf(file):
    fn = file.name
    if fn.endswith('.bib'):
        # We first have to convert the data in the bib file to a readable string format using the unicode transformation format - 8 bits; and eventually return a dataframe
        content = file.read().decode('utf-8')
        bib_db = bibtexparser.load(io.StringIO(content))
        return pd.DataFrame(bib_db.entries)
    elif fn.endswith('.xlsx'):
        # Excel files are relatively simpler to deal with; they only have to be read using the read_excel() function of the pandas library
        return pd.read_excel(file)
    else:
        st.error("Unsupported file format. Please upload a .bib or .xlsx file.")
        return None

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# Defining a function called getPub which takes the author name as an argument and returns a dataframe that consists of separate libraries for each entry for the author

# Notice that we're using the cache_data decorator of streamlit even here; for the same purpose!
@st.cache_data
def getPub(authorName):
    # Using the search_author() function of the scholarly library, we're going to web crawl for publications by authorname 
    query = scholarly.search_author(authorName)
    for _ in range(3):
        try:
            author = scholarly.fill(next(query))
            break
        except (StopIteration, MaxTriesExceededException):
            time.sleep(5)
    else:
        st.error("Failed to fetch data from Google Scholar. Try again later.")
        return pd.DataFrame()
    # It is a good practice to make a list (called pubs) and put dictionaries as its enteries (these dictionaries consist of information about one publication by the author, each) 
    pubs = [
        {
            'title': pub.get('bib', {}).get('title', ''),
            'year': pub.get('bib', {}).get('pub_year', ''),
            'journal': pub.get('bib', {}).get('venue', ''),
            'authors': pub.get('bib', {}).get('author', '')
        }
        for pub in author['publications']
    ]
    return pd.DataFrame(pubs)

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# Defining a function called filByYear that takes the data, the start year, and the end year as its arguments and returns the start and end points for our sliders

def filByYear(data, start, end):
    # Note that instead of using the conventional way of handling the ValueError using the try-except blocks, we chose to use the use the "errors" argument of the to_numeric() function
    # What this does is that it returns "NaN" in case an entry in the column "years" is, well, not a number, instead of trying to convert it anyway and then throwing a ValueError!
    data['year'] = pd.to_numeric(data['year'], errors='coerce')
    return data[(data['year'] >= start) & (data['year'] <= end)]

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# Defining a function called saveExcel which takes the data and the filename as arguments and enables the user to save the processed dataframe, so as to say, in excel format
def saveExcel(data, filename):
    excel_stream = BytesIO()

    with pd.ExcelWriter(excel_stream, engine='openpyxl') as writer:
        data.to_excel(writer, index=False, sheet_name='Sheet1')

    # The excel_stream.seek(0) method moves the file pointer to the beginning (position 0) of the file-like object  (here, the excel_stream).
    excel_stream.seek(0)

    st.download_button(
        label="Download Excel File",
        data=excel_stream,
        file_name=filename,
        # MIME stands for Multipurpose Internet Mail Extensions type and helps the system, browser, or application understand how to process or display the contents of a file.
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="download_button_excel_unique_key"
    )
    st.success(f"Data saved to {filename}. You can download it using the link above.")

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# Defining a function called saveWordl which takes the data and the filename as arguments and enables the user to save the processed dataframe, so as to say, in word format

def saveWord(data, filename):
    doc = Document()
    doc.add_heading('Publication Summary', 0)

    # A neat way of representing the data of a dataset is as follows. It iterates over the contents of the rows and presents them as seperate paragraphs in our word document
    for _, row in data.iterrows():
        doc.add_paragraph(f"{row['year']}: {row['title']} ({row['journal']})")

    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    st.download_button(
        label="Download Publication Summary",
        data=doc_stream,
        file_name=filename,
        # Note the difference in mime formats of the excel file above and the word doc below 
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_button_unique_key"
    )

    st.success("You can download the file using the link above.")

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# Now comes the main() function, where we shall call all the previous functions as and when necessary
def main():
    st.title("Faculty Publication Summary Tool")
    # Following is an illustration of how miraculous streamlit is; with one line of code, we can add an uploader that allows the user to use files from their own computer
    df = st.sidebar.file_uploader("Upload a BibTeX or Excel file", type=["bib", "xlsx"])

    if df:
        data = readdf(df)
        if data is not None:
            st.write("### Uploaded Data", data.head())
            if 'authorName' not in st.session_state:
                st.session_state.authorName = "Jane Doe" # A default value
            authorName = st.text_input("Enter Author Name", st.session_state.authorName)

            # In order to enhance the UX, we decided to ask the user what search option they want to go with. As we develop the idea more, we can add more such elements to make the UX seamless
            searchOpt = st.radio(
                "Choose search option:",
                ('Search within the uploaded dataset', 'Search universally using Google Scholar')
            )

            if st.button("Get Publications"):
                st.session_state.authorName = authorName
                if searchOpt == 'Search within the uploaded dataset':
                    publications = data[data['author'].str.contains(authorName, case=False, na=False)]
                    if not publications.empty:
                        st.session_state.publications = publications
                        st.write(f"### Publications for {authorName} (from uploaded dataset)", publications)
                    else:
                        st.warning(f"No publications found for {authorName} in the uploaded dataset")
                else:
                    with st.spinner('Fetching publications...'):
                        # Note that the getPub() function is only utlized when the user wishes to search for publications universally
                        publications = getPub(authorName)
                    if not publications.empty:
                        st.session_state.publications = publications
                        st.write(f"### Publications for {authorName} (from Google Scholar)", publications)
                    else:
                        st.warning(f"No publications found for {authorName}")

            if 'publications' in st.session_state:
                year1 = st.slider("Start Year", 1900, 2024, 2015)
                year2 = st.slider("End Year", 1900, 2024, 2020)
                filtered_data = filByYear(st.session_state.publications, year1, year2)
                st.write(f"### Filtered Publications ({year1}-{year2})", filtered_data)

                if st.button("Save to Excel"):
                    saveExcel(filtered_data, 'publication_summary.xlsx')
                if st.button("Save to Word"):
                    saveWord(filtered_data, 'publication_summary.docx')

# Ensures that the program functions only when the code is run directly, and not when it is imported as a module in another script; and calling main()
if __name__ == "__main__":
    main()
